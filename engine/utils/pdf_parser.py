
import io
from typing import Optional

async def extract_text_from_pdf(content: bytes, filename: str = "") -> Optional[str]:
    """
    Extract text from PDF or DOCX file bytes.   
    Args:
        content: Raw file bytes
        filename: Original filename (used to determine file type)    
    Returns:
        Extracted text or None if extraction fails
    """
    text = None
       
    if filename.lower().endswith(('.docx', '.doc')):
        text = await _extract_from_docx(content)
        if text:
            return text
    text = await _extract_from_pdf_pypdf(content)
    if text:
        return text
    text = await _extract_from_pdf_pdfplumber(content)
    if text:
        return text
    if not filename.lower().endswith(('.docx', '.doc')):
        text = await _extract_from_docx(content)    
    return text
async def _extract_from_pdf_pypdf(content: bytes) -> Optional[str]:
    """Extract text using pypdf (fast, works for most PDFs)."""
    try:
        import pypdf
        
        pdf_file = io.BytesIO(content)
        reader = pypdf.PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
        
        if text_parts:
            return "\n\n".join(text_parts)
        return None
        
    except ImportError:
        print("pypdf not installed, skipping...")
        return None
    except Exception as e:
        print(f"pypdf extraction failed: {e}")
        return None


async def _extract_from_pdf_pdfplumber(content: bytes) -> Optional[str]:
    """Extract text using pdfplumber (better for complex layouts)."""
    try:
        import pdfplumber
        
        pdf_file = io.BytesIO(content)
        text_parts = []
        
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
        
        if text_parts:
            return "\n\n".join(text_parts)
        return None
        
    except ImportError:
        print("pdfplumber not installed, skipping...")
        return None
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}")
        return None


async def _extract_from_docx(content: bytes) -> Optional[str]:
    """Extract text from DOCX files."""
    try:
        from docx import Document
        
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if text_parts:
            return "\n".join(text_parts)
        return None
        
    except ImportError:
        print("python-docx not installed, skipping...")
        return None
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return None


def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    """
    if not text:
        return ""
    
    import re
    
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    
    # Remove null characters
    text = text.replace('\x00', '')
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    return text.strip()
